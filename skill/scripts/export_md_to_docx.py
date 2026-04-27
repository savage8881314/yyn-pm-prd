import argparse
import html
import os
import shutil
import subprocess
import sys
import uuid
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: python-docx. Install it with `pip install python-docx` and rerun."
    ) from exc

try:
    from playwright.sync_api import sync_playwright
except ImportError:  # pragma: no cover
    sync_playwright = None


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(95, 99, 104)
TABLE_HEADER = "D9EAF7"
CODE_BG = "F3F5F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = ACCENT

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    idx = start
    while idx < len(lines) and is_table_line(lines[idx]):
        table_lines.append(lines[idx].strip())
        idx += 1

    rows = []
    for i, line in enumerate(table_lines):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if i == 1 and all(cell.replace("-", "").replace(":", "").strip() == "" for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.font.bold = True
            if r_idx == 0:
                set_cell_shading(cell, TABLE_HEADER)

    doc.add_paragraph("")


def add_code_block(doc: Document, code_lines: list[str], block_type: str = "") -> None:
    label = "Mermaid 图表源码" if block_type == "mermaid" else "代码块"
    caption = doc.add_paragraph()
    caption_run = caption.add_run(label)
    caption_run.font.name = "Microsoft YaHei"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption_run.font.size = Pt(9.5)
    caption_run.font.italic = True
    caption_run.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(4)

    paragraph = doc.add_paragraph()
    set_paragraph_shading(paragraph, CODE_BG)
    paragraph.paragraph_format.left_indent = Cm(0.2)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)

    for idx, line in enumerate(code_lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        if idx < len(code_lines) - 1:
            run.add_break()


def add_cover_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    if not pairs:
        return
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(pairs):
        left = table.cell(row_idx, 0)
        right = table.cell(row_idx, 1)
        left.text = label
        right.text = value
        set_cell_shading(left, TABLE_HEADER)
        for cell in (left, right):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)
        for run in left.paragraphs[0].runs:
            run.font.bold = True
    doc.add_paragraph("")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_mermaid_image(doc: Document, image_path: Path) -> None:
    caption = doc.add_paragraph()
    caption_run = caption.add_run("Mermaid 图表")
    caption_run.font.name = "Microsoft YaHei"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption_run.font.size = Pt(9.5)
    caption_run.font.italic = True
    caption_run.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(4)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(14.2))
    paragraph.paragraph_format.space_after = Pt(8)


def resolve_mermaid_script() -> str:
    env_path = os.environ.get("MERMAID_JS_PATH")
    if env_path:
        path = Path(env_path)
        if path.exists():
            return path.resolve().as_uri()

    candidates = [
        Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "node" / "node_modules" / "mermaid" / "dist" / "mermaid.min.js",
        Path.cwd() / "node_modules" / "mermaid" / "dist" / "mermaid.min.js",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate.resolve().as_uri()

    return "https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"


def render_with_external_python(html_path: Path, image_path: Path) -> bool:
    candidates = []
    env_python = os.environ.get("MERMAID_RENDER_PYTHON")
    if env_python:
        candidates.append(env_python)

    py_launcher = shutil.which("py")
    if py_launcher:
        candidates.append(f"{py_launcher} -3")

    system_python = shutil.which("python")
    if system_python:
        candidates.append(system_python)

    if sys.executable:
        candidates.append(sys.executable)

    inline_script = r"""
import sys
from pathlib import Path
from playwright.sync_api import sync_playwright

html_path = Path(sys.argv[1]).resolve()
image_path = Path(sys.argv[2]).resolve()
with sync_playwright() as playwright:
    browser = playwright.chromium.launch(headless=True)
    page = browser.new_page(viewport={"width": 2400, "height": 1800}, device_scale_factor=3)
    page.goto(html_path.as_uri(), wait_until="load", timeout=90000)
    page.wait_for_selector("svg", timeout=20000)
    page.locator(".canvas").first.screenshot(path=str(image_path))
    browser.close()
"""

    seen = set()
    for candidate in candidates:
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        try:
            if " " in candidate:
                cmd = candidate.split(" ")
            else:
                cmd = [candidate]
            completed = subprocess.run(
                cmd + ["-c", inline_script, str(html_path), str(image_path)],
                check=False,
                capture_output=True,
                text=True,
                timeout=120,
            )
            if completed.returncode == 0 and image_path.exists():
                return True
        except Exception:
            continue
    return False


def render_mermaid_to_image(code_lines: list[str], image_path: Path) -> bool:
    image_path.parent.mkdir(parents=True, exist_ok=True)
    mermaid_js = resolve_mermaid_script()
    graph_code = "\n".join(code_lines)
    html_content = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <style>
    body {{
      margin: 0;
      padding: 32px;
      background: white;
      font-family: "Segoe UI", Arial, sans-serif;
    }}
    .canvas {{
      display: inline-block;
      background: white;
      padding: 16px;
      min-width: 1800px;
    }}
    .mermaid {{
      font-size: 22px;
    }}
  </style>
</head>
<body>
  <div class="canvas">
    <div class="mermaid">
{html.escape(graph_code)}
    </div>
  </div>
  <script src="{mermaid_js}"></script>
  <script>
    mermaid.initialize({{ startOnLoad: true, securityLevel: "loose", theme: "default" }});
  </script>
</body>
</html>
"""

    temp_dir = image_path.parent / f"_render_tmp_{uuid.uuid4().hex}"
    temp_dir.mkdir(parents=True, exist_ok=True)
    html_path = temp_dir / "mermaid_render.html"
    try:
        html_path.write_text(html_content, encoding="utf-8")
        if sync_playwright is None:
            return render_with_external_python(html_path, image_path)
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(headless=True)
            page = browser.new_page(viewport={"width": 2400, "height": 1800}, device_scale_factor=3)
            page.goto(html_path.resolve().as_uri(), wait_until="load", timeout=90000)
            page.wait_for_selector("svg", timeout=20000)
            locator = page.locator(".canvas").first
            locator.screenshot(path=str(image_path))
            browser.close()
        return True
    except Exception:
        return False
    finally:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass


def extract_title_and_summary(lines: list[str]) -> tuple[str | None, list[tuple[str, str]], int]:
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    if idx >= len(lines) or not lines[idx].startswith("# "):
        return None, [], 0

    title = lines[idx][2:].strip()
    idx += 1
    summary_pairs: list[tuple[str, str]] = []

    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            idx += 1
            continue
        if stripped.startswith("## "):
            break
        if stripped.startswith("- ") and "：" in stripped:
            content = stripped[2:]
            label, value = content.split("：", 1)
            summary_pairs.append((label.strip(), value.strip()))
        else:
            break
        idx += 1

    return title, summary_pairs, idx


def convert_markdown(source: Path, target: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_style(doc)

    asset_dir = target.parent / f"{target.stem}_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)
    mermaid_count = 0

    title, summary_pairs, idx = extract_title_and_summary(lines)
    if title:
        title_paragraph = doc.add_paragraph(style="Title")
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.name = "Microsoft YaHei"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        add_cover_table(doc, summary_pairs)
        doc.add_paragraph("")
    else:
        idx = 0

    in_code = False
    code_lines: list[str] = []
    code_lang = ""

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                if code_lang == "mermaid":
                    mermaid_count += 1
                    image_path = asset_dir / f"mermaid_{mermaid_count}.png"
                    if render_mermaid_to_image(code_lines, image_path):
                        add_mermaid_image(doc, image_path)
                    else:
                        add_code_block(doc, code_lines, code_lang)
                else:
                    add_code_block(doc, code_lines, code_lang)
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                in_code = True
                code_lang = stripped[3:].strip().lower()
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if is_table_line(line):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(stripped[2:].strip())
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
        elif len(stripped) > 3 and stripped[0].isdigit() and stripped[1:3] == ". ":
            add_numbered(doc, stripped[3:].strip())
        else:
            add_body_paragraph(doc, line)
        idx += 1

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def main() -> int:
    parser = argparse.ArgumentParser(description="将 Markdown PRD 导出为更适合评审的 Word 文档。")
    parser.add_argument("--input", required=True, help="输入 Markdown 文件路径。")
    parser.add_argument("--output", required=True, help="输出 DOCX 文件路径。")
    args = parser.parse_args()

    convert_markdown(Path(args.input).resolve(), Path(args.output).resolve())
    print(f"Word 文档已生成: {Path(args.output).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
